#modules importing

from docx import Document
import os
import PyPDF2

#path and files declration 
path = os.getcwd()
pth = os.path.join(path,"media/")

docname = []
docslist = []
pdflist =[]
pdffilesname = [] 
#checking and seprating pdf and document file with file name and path
for files in os.listdir(pth):
    if files.endswith('.pdf'):
        str2 = files[:-4]
        pdffilesname.append(str2)
        pdflist.append(pth+files)
    elif files.endswith('.docx'):
        str2 = files[:-5]
        docname.append(str2)
        docslist.append(path+"/media/"+files)    
#pdf operations using pypdf2        
for pdfs in pdflist:
    pdfObj = PyPDF2.PdfFileReader(pdfs)
    pg = pdfObj.numPages
    str = ""
    for i in range(1,pg):
        str += (pdfObj.getPage(i).extractText())
#saving extracted text file in folder        
    for name in pdffilesname:
        text_file = open(os.path.join(path,"PDF_Extracted_Data/{}.txt".format(name)), "w")
        text_file.write(str)
        text_file.close()
#text extraction from docx using docx         
for doc in docslist:   
        doc = Document(doc)
        str = ''
        for para in doc.paragraphs:
            str += (para.text)             
#saving extracted text file in folder        
        for name in docname:
            text_file = open(os.path.join(path,"Doc_Extracted_Data/{}.txt".format(name)), "w")
            text_file.write(str)
            text_file.close()
print("Extraction DONE ! ")  











# # from xml.dom.minidom import Document
# # from docx import Document
# # import PyPDF2, os   
# # pdfFiles = []
# # path = "/home/akashk/Desktop/Workspace/project1/media"
# # for filename in os.listdir(path):
# #     if filename.endswith('.pdf'):
# #             pdfFiles.append(filename)
# #             pdfFiles.sort(key=str.lower)
# #             pdfWriter = PyPDF2.PdfFileWriter()
# #             for filename in pdfFiles:
# #                 # pdfFileObj = open(filename, 'rb')
# #                 pdfReader = PyPDF2.PdfFileReader(filename)
# #                 print(pdfReader)
#     #         for pageNum in range(0, pdfReader.numPages):
#     #                    pageObj = pdfReader.getPage(pageNum)
#     #                    text=pageObj.extractText() 
#     #                    file1=open("Extracted_text_pdf.txt",'a')
#     #                    file1.writelines(text)
#     #                    file1.close()    
#     # elif filename.endswith('.doc'):
#     #             doc = Document(filename)
#     #             text = []
#     #             for line in doc.paragraphs:
#     #                 text.append(line.text)
#     #             file1=open("Extracted_text_doc.txt",'a')
#     #             file1.writelines(text)
#     #             file1.close() 
#     # else:
#     #         text = "No text found"
#     #         file1=open("Balnk.txt",'a')
#     #         file1.writelines(text)
#     #         file1.close()            




# # string1 = 'A.pdf'
# # print(string1[:-4])
# from docx import Document
# import os 

# docname = []
# docslist = []
# path1 = os.getcwd()
# path = os.path.join(path1,'media/')
# for files in os.listdir(path):
#     if files.endswith('.docx'):
#         str2 = files[:-5]
#         docname.append(str2)
#         docslist.append(path+files)
        
# for doc in docslist:    
#         doc = Document(doc)
#         str = ''
#         for para in doc.paragraphs:
#             str += (para.text)
       
#         for name in docname:
#             text_file = open(os.path.join(path1,"Doc_Extracted_Data/{}.txt".format(name)), "w")
#             text_file.write(str)
#             text_file.close()
# print("Extraction DONE ! ")  